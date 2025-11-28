"""
Parser de documentos para extração de texto.
Autor: Matheus C. Pestana

Suporta: .txt, .docx, .doc, .pdf
"""

from pathlib import Path
from typing import Union
import io


def extract_text(source: Union[str, Path, bytes], file_type: str = None) -> str:
    """
    Extrai texto de um arquivo ou bytes.
    
    Args:
        source: Caminho do arquivo, Path object, ou bytes do arquivo
        file_type: Tipo do arquivo (txt, docx, doc, pdf). Auto-detecta se None.
    
    Returns:
        Texto extraído do documento
    
    Raises:
        ValueError: Se o tipo de arquivo não for suportado
        FileNotFoundError: Se o arquivo não existir
    """
    # Se for string de texto puro (não caminho de arquivo)
    if isinstance(source, str) and not Path(source).suffix:
        # Verifica se parece ser um caminho
        if not ('/' in source or '\\' in source or Path(source).exists()):
            return source  # É texto puro
    
    # Converte para Path se necessário
    if isinstance(source, str):
        source = Path(source)
    
    # Se for Path, lê o arquivo
    if isinstance(source, Path):
        if not source.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {source}")
        
        file_type = file_type or source.suffix.lower().lstrip('.')
        
        with open(source, 'rb') as f:
            file_bytes = f.read()
    else:
        file_bytes = source
        if not file_type:
            raise ValueError("file_type é obrigatório quando source é bytes")
    
    # Extrai texto baseado no tipo
    extractors = {
        'txt': _extract_txt,
        'docx': _extract_docx,
        'doc': _extract_doc,
        'pdf': _extract_pdf,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
    
    return extractor(file_bytes)


def _extract_txt(file_bytes: bytes) -> str:
    """Extrai texto de arquivo .txt."""
    # Tenta diferentes encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    
    # Fallback: ignora erros
    return file_bytes.decode('utf-8', errors='ignore')


def _extract_docx(file_bytes: bytes) -> str:
    """Extrai texto de arquivo .docx usando docx2txt (mais rápido) ou python-docx (fallback)."""
    import tempfile
    import os
    
    # Tenta docx2txt primeiro (mais rápido conforme benchmark)
    try:
        import docx2txt
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            text = docx2txt.process(tmp_path)
            return text
        finally:
            os.unlink(tmp_path)
    
    except ImportError:
        pass
    
    # Fallback para python-docx
    from docx import Document
    
    doc = Document(io.BytesIO(file_bytes))
    
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Também extrai texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragraphs.append(' | '.join(row_text))
    
    return '\n'.join(paragraphs)


def _extract_doc(file_bytes: bytes) -> str:
    """Extrai texto de arquivo .doc usando textract."""
    try:
        import textract
        import tempfile
        import os
        
        # textract precisa de um arquivo no disco
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            text = textract.process(tmp_path).decode('utf-8')
        finally:
            os.unlink(tmp_path)
        
        return text
    except ImportError:
        raise ImportError(
            "textract não está instalado. "
            "Instale com: uv pip install textract"
        )
    except Exception as e:
        # Fallback: tenta docx2txt (pode funcionar para alguns .doc)
        try:
            import docx2txt
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            try:
                text = docx2txt.process(tmp_path)
            finally:
                os.unlink(tmp_path)
            
            return text
        except:
            raise RuntimeError(
                f"Não foi possível extrair texto do .doc: {e}. "
                "Certifique-se de que antiword está instalado."
            )


def _extract_pdf(file_bytes: bytes) -> str:
    """Extrai texto de arquivo .pdf."""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n'.join(text_parts)
    except ImportError:
        try:
            # Fallback para PyPDF2
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
        except ImportError:
            raise ImportError(
                "pypdf ou PyPDF2 não está instalado. "
                "Instale com: uv pip install pypdf"
            )


def get_supported_extensions() -> list:
    """Retorna lista de extensões suportadas."""
    return ['txt', 'docx', 'doc', 'pdf']

